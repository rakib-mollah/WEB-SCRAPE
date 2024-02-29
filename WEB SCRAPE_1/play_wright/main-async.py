from playwright.async_api import async_playwright
from docx import Document
import asyncio
check = 1
p1 = 1
p2 = 1
p3 = 1
async def main():
    async with async_playwright() as p:
        # Launch the browser
        browser = await p.chromium.launch()

        # Create a new page
        page = await browser.new_page()

        try:
            # Navigate to the main webpage
            await page.goto("https://www.prothomalo.com/topic/%E0%A6%95%E0%A6%AC%E0%A6%BF%E0%A6%A4%E0%A6%BE-%E0%A6%85%E0%A6%A8%E0%A7%8D%E0%A6%AF-%E0%A6%86%E0%A6%B2%E0%A7%8B")

            # Wait for 40 seconds for the content to load
            await page.wait_for_timeout(40000)

            # Function to click the "Load More" button
            async def click_load_more():
                try:
                    await page.click(".tNj8k")
                    return True
                except:
                    return False

            # Click the "Load More" button until it's no longer available
            cnt = 1
            while cnt <= 1:
                await click_load_more()
                await asyncio.sleep(50)  # Wait for content to load
                cnt += 1

            # Find all anchor elements by XPath
            links = await page.query_selector_all(".title-link")

            # Extract href attributes from anchor elements
            hrefs = [await link.get_attribute("href") for link in links]

            # Create a new Word document for hrefs
            doc_hrefs = Document()

            # Save hrefs to the document
            for href in hrefs:
                doc_hrefs.add_paragraph(href)

            # Save the document with hrefs
            doc_hrefs.save("extracted_hrefs.docx")

            print("Hrefs extracted and saved to 'extracted_hrefs.docx'")

            # Create Word documents for extracted text, writer, poem, and full content
            doc_text = Document()
            doc_writer = Document()
            doc_poem = Document()
            doc_full = Document()

            number = 1
            for href in hrefs:
                # Visit each link
                await page.goto(href)
                await asyncio.sleep(5)  # Wait for 5 seconds to ensure page loads completely

                # Find element by class name and extract text
                element_text = (await page.query_selector(".IiRps"))

                # print("ok")
                global p1; print("P1 : " + str(p1));p1 += 1

                # Find writer element
                element_writer = (await page.query_selector(".mH-nJ"))

                # print("ok")
                global p2; print("P2 : " + str(p2));p2 += 1

                # Find poem elements
                elements_poem = await page.query_selector_all(".paPvR p")

                # print("ok")
                global p3; print("P3 : " + str(p3));p3 += 1

                text_content = await element_text.text_content() if element_text else None
                writer_content = await element_writer.text_content() if element_writer else None

                texts = []
                for all_text in elements_poem:
                    text = await all_text.text_content()
                    if text:
                        texts.append(text)

                poem_content = '\n'.join(texts)

                # print("ok : "+ str(check++))
                global check; print("ok : " + str(check));check += 1


                # Add the extracted text to the respective documents
                doc_text.add_paragraph(text_content)
                doc_writer.add_paragraph(writer_content)
                doc_poem.add_paragraph(poem_content)

                # Add the full content to the document
                # result = "\n".join([str(number) + ".................", element_text, element_writer + "\n", element_poem + "\n\n"])
                # result = "\n".join([
                #     str(number) + ".................", 
                #     await text_content.text_content() if hasattr(text_content, 'text_content') else "", 
                #     (await writer_content.text_content() if hasattr(writer_content, 'text_content') else "") + "\n", 
                #     poem_content + "\n\n" if poem_content else ""
                # ])

                result = "\n".join([
                    str(number) + ".................", 
                    text_content if text_content else "", 
                    writer_content + "\n" if writer_content else "", 
                    poem_content + "\n\n" if poem_content else ""
                ])


                number += 1
                doc_full.add_paragraph(result)

            # Save the documents with extracted text, writer, poem, and full content
            doc_text.save("extracted_text.docx")
            doc_writer.save("extracted_writer.docx")
            doc_poem.save("extracted_poem.docx")
            doc_full.save("last.docx")

            print("Text extracted and saved to 'extracted_text.docx'")

        except Exception as e:
            print('An error occurred: ', str(e))
        finally:
            await browser.close()

asyncio.run(main())
