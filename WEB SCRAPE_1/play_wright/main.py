from playwright.sync_api import sync_playwright
from docx import Document
import time

# Initialize Playwright
with sync_playwright() as p:
    # Launch the browser
    browser = p.chromium.launch()

    # Create a new page
    page = browser.new_page()

    try:
        # Navigate to the main webpage
        page.goto("https://www.prothomalo.com/topic/%E0%A6%95%E0%A6%AC%E0%A6%BF%E0%A6%A4%E0%A6%BE-%E0%A6%85%E0%A6%A8%E0%A7%8D%E0%A6%AF-%E0%A6%86%E0%A6%B2%E0%A7%8B")

        # Wait for 40 seconds for the content to load
        page.wait_for_timeout(40000)

        # Function to click the "Load More" button
        def click_load_more():
            try:
                page.click(".tNj8k")
                return True
            except:
                return False

        # Click the "Load More" button until it's no longer available
        cnt = 1
        while cnt <= 5:
            click_load_more()
            time.sleep(50)  # Wait for content to load
            cnt += 1

        # Find all anchor elements by XPath
        links = page.query_selector_all(".title-link")

        # Extract href attributes from anchor elements
        hrefs = [link.get_attribute("href") for link in links]

        # Create a new Word document for hrefs
        doc_hrefs = Document()

        # Save hrefs to the document
        for href in hrefs:
            doc_hrefs.add_paragraph(href)

        # Save the document with hrefs
        doc_hrefs.save("extracted_hrefs.docx")

        print("Hrefs extracted and saved to 'extracted_hrefs.docx'")

        # Create Word documents for extracted text, writer, poem, and full content
        doc_text = Document()
        doc_writer = Document()
        doc_poem = Document()
        doc_full = Document()

        number = 1
        for href in hrefs:
            # Visit each link
            page.goto(href)
            time.sleep(5)  # Wait for 5 seconds to ensure page loads completely

            # Find element by class name and extract text
            element_text = page.query_selector(".IiRps").text_content()

            # Find writer element
            element_writer = page.query_selector(".mH-nJ").text_content()

            # Find poem elements
            elements_poem = page.query_selector_all(".paPvR p")

            texts = []
            for all_text in elements_poem:
                text = all_text.text_content()
                if text:
                    texts.append(text)

            element_poem = '\n'.join(texts)

            # Add the extracted text to the respective documents
            doc_text.add_paragraph(element_text)
            doc_writer.add_paragraph(element_writer)
            doc_poem.add_paragraph(element_poem)

            # Add the full content to the document
            result = "\n".join([str(number) + ".................", element_text, element_writer + "\n", element_poem + "\n\n"])
            number += 1
            doc_full.add_paragraph(result)

        # Save the documents with extracted text, writer, poem, and full content
        doc_text.save("extracted_text.docx")
        doc_writer.save("extracted_writer.docx")
        doc_poem.save("extracted_poem.docx")
        doc_full.save("last.docx")

        print("Text extracted and saved to 'extracted_text.docx'")

    except Exception as e:
        print('An error occurred: ', str(e))
    finally:
        browser.close()
