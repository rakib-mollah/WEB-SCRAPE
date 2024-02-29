from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.support.wait import WebDriverWait
from selenium.common.exceptions import WebDriverException



from docx import Document
import time

# Create Chrome options
chrome_options = Options()
chrome_options.add_argument('--lang=en_US.UTF-8')

# chrome_options.add_argument("--headless")  # Optional: run headless to hide the browser window

# Initialize Chrome driver with the proper executable path
driver = webdriver.Chrome(options=chrome_options)

try:
    # Navigate to the main webpage
    driver.get("https://www.prothomalo.com/topic/%E0%A6%95%E0%A6%AC%E0%A6%BF%E0%A6%A4%E0%A6%BE-%E0%A6%85%E0%A6%A8%E0%A7%8D%E0%A6%AF-%E0%A6%86%E0%A6%B2%E0%A7%8B")
    
    # Wait for 5 seconds
    # time.sleep(15)
    wait = WebDriverWait(driver, 40)

    
    # Function to click the "Load More" button
    def click_load_more():
        try:
            load_more_button = driver.find_element(By.CLASS_NAME, "tNj8k")
            load_more_button.click()
            return True
        except WebDriverException as e:
            print(e)
            return False
    
    # Click the "Load More" button until it's no longer available
    cnt=1
    while cnt<=5:
        click_load_more()
        time.sleep(50)  # Wait for content to load
        cnt+=1
    
    # Find all anchor elements by XPath
    links = driver.find_elements(By.CLASS_NAME, "title-link")
    
    # Extract href attributes from anchor elements
    hrefs = [link.get_attribute("href") for link in links]
    
    # Create a new Word document for hrefs
    doc_hrefs = Document()
    
    # Save hrefs to the document
    for href in hrefs:
        doc_hrefs.add_paragraph(href)
    
    # Save the document with hrefs
    doc_hrefs.save("extracted_hrefs.docx")
    
    print("Hrefs extracted and saved to 'extracted_hrefs.docx'")
    
    # Create a new Word document for extracted text print-authors-list zM7oX
    doc_text = Document()
    doc_writer = Document()
    doc_poem = Document()
    doc_full = Document()
    # cnt = 1
    number = 1
    for href in hrefs:
        # Visit each link
        driver.get(href)
        time.sleep(5)  # Wait for 5 seconds to ensure page loads completely
        
        # Find element by class name and extract text
        try:
            element_text = driver.find_element(By.CLASS_NAME, "IiRps").get_attribute("textContent")
        except:
            element_text = "Text not found"  # Handle cases where the class is not found or text is not available

        try:
            element_writer = driver.find_element(By.CLASS_NAME, "mH-nJ").get_attribute("textContent")
        except:
            element_writer = "Text not found" 

        try:
            # element_poem = driver.find_element(By.CLASS_NAME, "storyCard eyOoS").get_attribute("textContent")
            division_element = driver.find_element(By.CLASS_NAME, "paPvR ")
            elements_poem = division_element.find_elements(By.TAG_NAME, "p")
        except NoSuchElementException:
            elements_poem = []


        texts = []
        for all_text in elements_poem:
            # all_text.get_attribute("textContent")
            text = all_text.get_attribute("textContent")
            if text:
                texts.append(text)


        element_poem = '\n'.join(texts)   
        # Add the extracted text to the document
        doc_text.add_paragraph(element_text)
        doc_writer.add_paragraph(element_writer)
        doc_poem.add_paragraph(element_poem)
        
        result = "\n".join([str(number)+".................",element_text, element_writer+"\n", element_poem+"\n\n"])
        number += 1
        doc_full.add_paragraph(result)
        # cnt += 1
        
        # if cnt == 3:  # Limit to 3 for demonstration, you can change this limit or remove it
        #     break
    
    # Save the document with extracted text
    doc_text.save("extracted_text.docx")
    doc_writer.save("extracted_writer.docx")
    doc_poem.save("extracted_poem.docx")
    doc_full.save("last.docx")
    
    print("Text extracted and saved to 'extracted_text.docx'")
    
except WebDriverException as e:
    print('An error occurred: ', str(e))
finally:
    driver.quit()
