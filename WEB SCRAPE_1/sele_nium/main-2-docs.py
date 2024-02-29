from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from docx import Document
import time

# Create Chrome options
chrome_options = Options()
chrome_options.add_argument('--lang=en_US.UTF-8')

# chrome_options.add_argument("--headless")  # Optional: run headless to hide the browser window

# Initialize Chrome driver with the proper executable path
driver = webdriver.Chrome(options=chrome_options)


try:
    # Navigate to the main webpage
    driver.get("https://www.prothomalo.com/topic/%E0%A6%95%E0%A6%AC%E0%A6%BF%E0%A6%A4%E0%A6%BE-%E0%A6%85%E0%A6%A8%E0%A7%8D%E0%A6%AF-%E0%A6%86%E0%A6%B2%E0%A7%8B")
    
    # Wait for 5 seconds
    time.sleep(2)
    
    # Find all anchor elements by XPath
    links = driver.find_elements(By.CLASS_NAME, "title-link")

    
    # Extract href attributes from anchor elements
    hrefs = [link.get_attribute("href") for link in links]
    
    # Create a new Word document for hrefs
    doc_hrefs = Document()
    
    # Save hrefs to the document
    for href in hrefs:
        doc_hrefs.add_paragraph(href)
    
    # Save the document with hrefs
    doc_hrefs.save("extracted_hrefs.docx")
    
    print("Hrefs extracted and saved to 'extracted_hrefs.docx'")
    
    # Create a new Word document for extracted text
    # author-location _8-umj
    doc_text = Document()
    cnt=1
    for href in hrefs:
        # Visit each link
        driver.get(href)
        time.sleep(5)  # Wait for 3 seconds to ensure page loads completely
        
        # Find element by class name and extract text
        try:
            element_text = driver.find_element(By.CLASS_NAME, "IiRps")
            element_text = element_text.get_attribute("textContent")
        except:
            element_text = "Text not found"  # Handle cases where the class is not found or text is not available
            
        # Add the extracted text to the document
        doc_text.add_paragraph(element_text)
        # cnt=cnt+1
        if(cnt==3):
            break
    
    # Save the document with extracted text
    doc_text.save("extracted_text.docx")
    
    print("Text extracted and saved to 'extracted_text.docx'")
    
except Exception as e:
    print('An error occurred: ', str(e))
finally:
    driver.quit()
